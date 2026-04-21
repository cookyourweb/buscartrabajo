"""
Genera un CV adaptado en .docx y lo sube a Google Drive.
N8N lo llama con: python3 generar_cv_drive.py "Empresa" "Puesto" "Texto del CV adaptado"

Instalación:
  pip install google-auth-oauthlib google-api-python-client python-docx

Uso desde N8N (Execute Command):
  cd /Users/vero/Desktop/buscartrabajo && source venv/bin/activate && python3 generar_cv_drive.py "{{ empresa }}" "{{ puesto }}" "{{ cv_texto }}"
"""

import sys
import os
import re
import json
from datetime import date
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pickle

# ── PARÁMETROS DESDE N8N ──────────────────────────────────────────────────────
EMPRESA  = sys.argv[1] if len(sys.argv) > 1 else "Empresa"
PUESTO   = sys.argv[2] if len(sys.argv) > 2 else "Puesto"
CV_TEXTO = sys.argv[3] if len(sys.argv) > 3 else ""
# ──────────────────────────────────────────────────────────────────────────────

FECHA       = date.today().strftime("%Y-%m-%d")
EMPRESA_SLUG = re.sub(r'[^a-zA-Z0-9]', '', EMPRESA)
FOLDER_NAME  = f"{EMPRESA_SLUG}_{FECHA}"
CV_FILENAME  = f"cv-veronicaSerna-{re.sub(r'[^a-zA-Z0-9]', '', PUESTO)[:20]}.docx"
CV_PATH      = f"/tmp/{CV_FILENAME}"

SCOPES    = ["https://www.googleapis.com/auth/drive"]
MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TOKEN_PATH = os.path.join(os.path.dirname(__file__), "token.pickle")
CREDS_PATH = os.path.join(os.path.dirname(__file__), "credentials.json")


def generar_docx(texto, empresa, puesto, output_path):
    """Genera un docx limpio con el CV adaptado."""
    doc = Document()

    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("Verónica Serna Pérez")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitulo.add_run("Tech Lead UX Engineer · AI & Automation Specialist")
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x1F, 0x5C, 0x8B)

    contacto = doc.add_paragraph()
    contacto.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = contacto.add_run("Madrid · verserper@gmail.com · linkedin.com/in/veronica4web")
    run3.font.size = Pt(9)
    run3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    # Empresa y puesto adaptado
    header = doc.add_paragraph()
    run4 = header.add_run(f"CV adaptado para: {empresa} — {puesto}")
    run4.bold = True
    run4.font.size = Pt(10)
    run4.font.color.rgb = RGBColor(0x1F, 0x5C, 0x8B)

    doc.add_paragraph()

    # Contenido del CV adaptado
    for linea in texto.strip().split('\n'):
        linea = linea.strip()
        if not linea:
            doc.add_paragraph()
            continue
        # Limpiar markdown
        linea = re.sub(r'\*\*(.*?)\*\*', r'\1', linea)
        linea = re.sub(r'^#{1,3}\s*', '', linea)
        linea = re.sub(r'^---+$', '', linea)
        if linea:
            p = doc.add_paragraph(linea)
            p.style.font.size = Pt(10)

    doc.save(output_path)


def autenticar():
    creds = None
    if os.path.exists(TOKEN_PATH):
        with open(TOKEN_PATH, "rb") as f:
            creds = pickle.load(f)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(CREDS_PATH, SCOPES)
            creds = flow.run_local_server(port=0)
        with open(TOKEN_PATH, "wb") as f:
            pickle.dump(creds, f)
    return build("drive", "v3", credentials=creds)


def crear_carpeta(service, nombre):
    res = service.files().list(
        q=f"name='{nombre}' and mimeType='application/vnd.google-apps.folder' and trashed=false",
        fields="files(id, name)"
    ).execute()
    if res["files"]:
        return res["files"][0]["id"]
    meta = {"name": nombre, "mimeType": "application/vnd.google-apps.folder"}
    folder = service.files().create(body=meta, fields="id").execute()
    return folder["id"]


def subir_cv(service, folder_id, cv_path, cv_filename):
    meta = {"name": cv_filename, "parents": [folder_id]}
    media = MediaFileUpload(cv_path, mimetype=MIME_DOCX, resumable=True)
    archivo = service.files().create(body=meta, media_body=media, fields="id, webViewLink").execute()
    return archivo.get("webViewLink")


if __name__ == "__main__":
    try:
        # 1. Generar docx
        generar_docx(CV_TEXTO, EMPRESA, PUESTO, CV_PATH)

        # 2. Subir a Drive
        service = autenticar()
        folder_id = crear_carpeta(service, FOLDER_NAME)
        link = subir_cv(service, folder_id, CV_PATH, CV_FILENAME)

        # 3. Output JSON para N8N
        result = {
            "success": True,
            "link": link,
            "carpeta": FOLDER_NAME,
            "archivo": CV_FILENAME
        }
        print(json.dumps(result))

    except Exception as e:
        print(json.dumps({"success": False, "error": str(e)}))
        sys.exit(1)
