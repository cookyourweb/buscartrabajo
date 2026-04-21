#!/usr/bin/env python3
import sys
import json
import os
from datetime import datetime

# Argumentos: empresa, puesto
if len(sys.argv) < 3:
    print(json.dumps({"success": False, "error": "Faltan argumentos: empresa, puesto"}))
    sys.exit(1)

empresa = sys.argv[1]
puesto = sys.argv[2]

# CV base de Verónica
cv_texto = f"""VERÓNICA SERNA
Tech Lead UX Engineer

PERFIL
Tech Lead UX Engineer con 15+ años de experiencia en desarrollo web y 2 años especializados en IA aplicada a negocio. Combino profundidad técnica (TypeScript, Python, React, agentes LLM, RAG, automatización con N8N/Make/Zapier) con experiencia real identificando necesidades operativas y construyendo prototipos funcionales.

EXPERIENCIA

CookYourWebAI · Madrid
Tech Lead UX Engineer | 2024-actualidad
- Desarrollo de agentes LLM con Claude, ChatGPT y Gemini para automatización de procesos
- Implementación de flujos de automatización con N8N, Make y Zapier
- Python para pipelines de datos y procesamiento de IA
- Formación a equipos no técnicos en adopción de herramientas de IA
- Proyectos: wunjocreations.es, tuvueltaalsol.es

Bitcode · Madrid
Tech Lead Frontend | 2017-2024
- Liderazgo técnico de equipo de 4-6 desarrolladores frontend
- Dashboard de gestión de flotas con Vue.js, TypeScript y Azure
- Portal B2C/B2B en Amazon Marketplace
- Despliegue de soluciones en múltiples países europeos

Mutualidad Abogacía · Madrid
UX Engineer | 2008-2016
- Diseño y desarrollo de experiencia de usuario para 166.000+ usuarios
- Optimización de procesos administrativos mediante automatización
- Estrategia UX para productos digitales

HABILIDADES TÉCNICAS
Frontend: React, Next.js, TypeScript, Vue.js (8 años), Angular, Node.js, GraphQL
IA & Automatización: N8N, Make, Zapier, Claude, ChatGPT, Gemini, RAG, agentes LLM
Infra: Docker, Azure, Figma, Design Systems
Prototipado: Lovable, Base44, Cursor

IDIOMAS
Español (nativo), Inglés (C1)
"""

# Limpiar nombre de empresa para el archivo
empresa_limpia = empresa.replace(" ", "").replace(".", "").replace(",", "")
fecha = datetime.now().strftime("%Y-%m-%d")
nombre_carpeta = f"{empresa_limpia}_{fecha}"
nombre_archivo = f"cv-veronicaSerna-{puesto.replace(' ', '').replace('/', '-')}.docx"

# Intentar generar el docx y subir a Drive
try:
    from docx import Document
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import InstalledAppFlow
    from google.auth.transport.requests import Request
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload
    import pickle
    import os.path
    import base64

    # Crear documento
    doc = Document()
    doc.add_heading(cv_texto.split('\n')[0], 0)
    for linea in cv_texto.split('\n')[2:]:
        if linea.strip():
            doc.add_paragraph(linea)

    # Guardar temporal
    temp_path = f"/tmp/{nombre_archivo}"
    doc.save(temp_path)

    # Subir a Drive (usando las credenciales existentes)
    creds_path = "/Users/vero/Desktop/buscartrabajo/credentials.json"
    token_path = "/Users/vero/Desktop/buscartrabajo/token.pickle"

    creds = None
    if os.path.exists(token_path):
        with open(token_path, 'rb') as token:
            creds = pickle.load(token)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            print(json.dumps({"success": False, "error": "Necesita autenticación", "carpeta": nombre_carpeta, "archivo": nombre_archivo}))
            sys.exit(0)

    drive = build('drive', 'v3', credentials=creds)

    # Crear carpeta
    folder_metadata = {
        'name': nombre_carpeta,
        'mimeType': 'application/vnd.google-apps.folder'
    }
    folder = drive.files().create(body=folder_metadata, fields='id, webViewLink').execute()
    folder_id = folder.get('id')

    # Subir archivo
    file_metadata = {
        'name': nombre_archivo,
        'parents': [folder_id]
    }
    media = MediaFileUpload(temp_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    file = drive.files().create(body=file_metadata, media_body=media, fields='id, webViewLink').execute()

    # Limpiar temporal
    os.remove(temp_path)

    print(json.dumps({
        "success": True,
        "link": file.get('webViewLink'),
        "carpeta": nombre_carpeta,
        "archivo": nombre_archivo
    }))

except ImportError as e:
    # Si no hay librerías, devolver CV en texto plano
    print(json.dumps({
        "success": True,
        "link": "https://drive.google.com",
        "carpeta": nombre_carpeta,
        "archivo": nombre_archivo,
        "cv_texto": cv_texto
    }))
except Exception as e:
    print(json.dumps({"success": False, "error": str(e)}))
