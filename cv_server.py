#!/usr/bin/env python3
"""
Servidor HTTP local para generar CVs y subirlos a Drive.
N8N llama a este servidor via HTTP Request en lugar de Execute Command.

Uso: python3 cv_server.py
Endpoint: POST http://localhost:8080/generar-cv
Body: {"empresa": "...", "puesto": "...", "cv_texto": "..."}
"""

from http.server import HTTPServer, BaseHTTPRequestHandler
import json
import sys
import os
from datetime import datetime
from docx import Document
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
import pickle

DIR_BASE = "/Users/vero/Desktop/buscartrabajo"

class CVHandler(BaseHTTPRequestHandler):
    def do_POST(self):
        if self.path != '/generar-cv':
            self.send_response(404)
            self.end_headers()
            return

        content_length = int(self.headers['Content-Length'])
        body = self.rfile.read(content_length)
        data = json.loads(body)

        empresa = data.get('empresa', 'Empresa')
        puesto = data.get('puesto', 'Puesto')
        cv_texto = data.get('cv_texto', '')

        resultado = generar_y_subir_cv(empresa, puesto, cv_texto)

        self.send_response(200)
        self.send_header('Content-Type', 'application/json')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        self.wfile.write(json.dumps(resultado).encode())

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()

def generar_y_subir_cv(empresa, puesto, cv_texto):
    try:
        # Limpiar nombres para archivo
        empresa_limpia = empresa.replace(" ", "").replace(".", "").replace(",", "").replace("/", "-")
        puesto_limpio = puesto.replace(" ", "").replace("/", "-").replace("(", "").replace(")", "")
        fecha = datetime.now().strftime("%Y-%m-%d")
        nombre_carpeta = f"{empresa_limpia}_{fecha}"
        nombre_archivo = f"cv-veronicaSerna-{puesto_limpio}.docx"

        # Crear documento
        doc = Document()

        # Si hay cv_texto usarlo, si no usar el CV base
        if cv_texto and len(cv_texto) > 50:
            lineas = cv_texto.split('\n')
            if lineas:
                doc.add_heading(lineas[0] if len(lineas[0]) < 100 else "VERÓNICA SERNA", 0)
            for linea in lineas[1:]:
                if linea.strip():
                    doc.add_paragraph(linea)
        else:
            # CV base por defecto
            cv_base = """VERÓNICA SERNA
Tech Lead UX Engineer

PERFIL
Tech Lead UX Engineer con 15+ años de experiencia en desarrollo web y 2 años especializados en IA aplicada a negocio.

EXPERIENCIA
CookYourWebAI · Madrid
Tech Lead UX Engineer | 2024-actualidad
- Agentes LLM con Claude, ChatGPT, Gemini
- Automatización N8N, Make, Zapier
- Python para pipelines de IA

Bitcode · Madrid
Tech Lead Frontend | 2017-2024
- Liderazgo equipo 4-6 devs
- Vue.js, TypeScript, Azure

HABILIDADES
React, TypeScript, Vue.js, Python, N8N, IA
"""
            for linea in cv_base.split('\n'):
                if linea.strip():
                    doc.add_paragraph(linea)

        # Guardar temporal
        temp_path = os.path.join(DIR_BASE, f"cv/{nombre_archivo}")
        os.makedirs(os.path.join(DIR_BASE, "cv"), exist_ok=True)
        doc.save(temp_path)

        # Credenciales Google
        creds_path = os.path.join(DIR_BASE, "credentials.json")
        token_path = os.path.join(DIR_BASE, "token.pickle")

        creds = None
        if os.path.exists(token_path):
            with open(token_path, 'rb') as token:
                creds = pickle.load(token)

        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                return {"success": False, "error": "Necesita autenticación Google"}

        drive = build('drive', 'v3', credentials=creds)

        # Crear carpeta
        folder_metadata = {'name': nombre_carpeta, 'mimeType': 'application/vnd.google-apps.folder'}
        folder = drive.files().create(body=folder_metadata, fields='id, webViewLink').execute()
        folder_id = folder.get('id')

        # Subir archivo
        file_metadata = {'name': nombre_archivo, 'parents': [folder_id]}
        media = MediaFileUpload(temp_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        file = drive.files().create(body=file_metadata, media_body=media, fields='id, webViewLink').execute()

        # Limpiar
        os.remove(temp_path)

        return {
            "success": True,
            "link": file.get('webViewLink'),
            "carpeta": nombre_carpeta,
            "archivo": nombre_archivo
        }

    except Exception as e:
        return {"success": False, "error": str(e)}

if __name__ == '__main__':
    port = 8080
    server = HTTPServer(('0.0.0.0', port), CVHandler)
    print(f"🚀 Servidor CV corriendo en http://0.0.0.0:{port} (accesible desde Docker via host.docker.internal:{port})")
    print(f"   Endpoint: POST /generar-cv")
    print(f"   Body: {{\"empresa\": \"...\", \"puesto\": \"...\", \"cv_texto\": \"...\"}}")
    server.serve_forever()
